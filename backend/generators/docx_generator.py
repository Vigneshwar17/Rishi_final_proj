"""
DOCX Generator using python-docx
Generates IEEE/Springer/ACM-styled DOCX with:
  - Title, Authors (single column front matter)
  - Two-column body via XML (w:cols)
  - Sections, References
"""

import logging
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


TEMPLATE_CONFIGS = {
    "ieee": {
        "title_size": 20,
        "author_size": 11,
        "heading_size": 11,
        "body_size": 10,
        "font_name": "Times New Roman",
        "margins": (Inches(0.75), Inches(0.75), Inches(1), Inches(1)),
    },
    "springer": {
        "title_size": 22,
        "author_size": 12,
        "heading_size": 12,
        "body_size": 10,
        "font_name": "Arial",
        "margins": (Inches(1), Inches(1), Inches(1), Inches(1)),
    },
    "acm": {
        "title_size": 18,
        "author_size": 11,
        "heading_size": 10,
        "body_size": 9,
        "font_name": "Arial",
        "margins": (Inches(0.75), Inches(0.75), Inches(1), Inches(1)),
    },
}


def generate_docx(doc_data, output_path: str, template: str = "ieee", styling: dict = None) -> str:
    styling = styling or {}
    cfg = TEMPLATE_CONFIGS.get(template.lower(), TEMPLATE_CONFIGS["ieee"])

    font_name = styling.get("fontFamily", cfg["font_name"])
    title_size = int(styling.get("titleSize", cfg["title_size"]))
    body_size = int(styling.get("bodySize", cfg["body_size"]))
    line_spacing = float(styling.get("lineSpacing", 1.15))

    doc = Document()

    # ── Page margins ──────────────────────────────────────
    section = doc.sections[0]
    L, R, T, B = cfg["margins"]
    section.left_margin = L
    section.right_margin = R
    section.top_margin = T
    section.bottom_margin = B

    # ── Title ─────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(doc_data.title or "Untitled Document")
    run.bold = True
    run.font.size = Pt(title_size)
    run.font.name = font_name
    _set_line_spacing(title_para, line_spacing)

    # ── Authors ───────────────────────────────────────────
    if doc_data.authors:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_texts = []
        for author in doc_data.authors:
            parts = []
            if author.name:
                parts.append(author.name)
            if author.institution:
                parts.append(author.institution)
            if author.email:
                parts.append(author.email)
            author_texts.append(", ".join(parts))
        run = author_para.add_run(" | ".join(author_texts))
        run.italic = True
        run.font.size = Pt(cfg["author_size"])
        run.font.name = font_name

    # ── Abstract ──────────────────────────────────────────
    if doc_data.abstract:
        _add_heading(doc, "Abstract", cfg["heading_size"], font_name)
        p = doc.add_paragraph(doc_data.abstract)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_font(p, font_name, body_size)
        _set_line_spacing(p, line_spacing)

    # ── Keywords ─────────────────────────────────────────
    if doc_data.keywords:
        kw_para = doc.add_paragraph()
        kw_run = kw_para.add_run("Keywords: ")
        kw_run.bold = True
        kw_run.font.name = font_name
        kw_run.font.size = Pt(body_size)
        kw_para.add_run(", ".join(doc_data.keywords)).font.name = font_name

    # ── Warnings ─────────────────────────────────────────
    if doc_data.missing_sections:
        warn_para = doc.add_paragraph()
        warn_run = warn_para.add_run("⚠ Missing: " + ", ".join(doc_data.missing_sections))
        warn_run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        warn_run.font.name = font_name
        warn_run.font.size = Pt(body_size)

    # ── Switch to 2-column layout ─────────────────────────
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.left_margin = L
    new_section.right_margin = R
    new_section.top_margin = T
    new_section.bottom_margin = B
    _set_two_columns(new_section)

    # ── Sections ──────────────────────────────────────────
    for section_data in doc_data.sections:
        _add_heading(doc, section_data.heading, cfg["heading_size"], font_name)
        for para_text in section_data.paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _set_font(p, font_name, body_size)
                _set_line_spacing(p, line_spacing)
        for eq in section_data.equations:
            eq_para = doc.add_paragraph(eq)
            eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_font(eq_para, font_name, body_size)

    # ── References ────────────────────────────────────────
    if doc_data.references:
        _add_heading(doc, "References", cfg["heading_size"], font_name)
        for i, ref in enumerate(doc_data.references, 1):
            p = doc.add_paragraph(f"[{i}] {ref}")
            _set_font(p, font_name, max(body_size - 1, 7))
            _set_line_spacing(p, 1.0)

    doc.save(output_path)
    logger.info(f"DOCX generated: {output_path}")
    return output_path


# ─────────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────────

def _add_heading(doc, text: str, size: int, font_name: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = font_name
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)


def _set_font(paragraph, font_name: str, size: int):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
    # Also set default run font for new text
    if not paragraph.runs:
        run = paragraph.add_run()
        run.font.name = font_name
        run.font.size = Pt(size)


def _set_line_spacing(paragraph, spacing: float):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing


def _set_two_columns(section):
    """
    Inject w:cols element into section properties for a two-column layout.
    """
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")  # 0.25 inch gap in twentieths of a point
    cols.set(qn("w:equalWidth"), "1")
