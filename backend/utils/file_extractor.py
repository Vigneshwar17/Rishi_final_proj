"""
File extraction utilities.
Handles: PDF, DOCX, TXT, MD
Extracts: raw text + embedded images (saved to /images)
"""

import os
import io
import re
import uuid
import logging

logger = logging.getLogger(__name__)

IMAGES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "images")
os.makedirs(IMAGES_DIR, exist_ok=True)


def extract_from_file(file_storage) -> dict:
    """
    Given a Werkzeug FileStorage object, extract text and images.
    Returns dict: { 'text': str, 'images': [str paths], 'error': str|None }
    """
    filename = file_storage.filename.lower() if file_storage.filename else ""
    content = file_storage.read()

    result = {"text": "", "images": [], "error": None}

    try:
        if filename.endswith(".pdf"):
            result = _extract_pdf(content)
        elif filename.endswith(".docx"):
            result = _extract_docx(content)
        elif filename.endswith(".txt") or filename.endswith(".md"):
            result["text"] = content.decode("utf-8", errors="replace")
        else:
            # Try as plain text as last resort
            result["text"] = content.decode("utf-8", errors="replace")
    except Exception as e:
        logger.error(f"Extraction error for {filename}: {e}")
        result["error"] = str(e)

    return result


def extract_from_text(raw_text: str) -> dict:
    """Wrap raw pasted text in the same return format."""
    return {"text": raw_text, "images": [], "error": None}


def _extract_pdf(content: bytes) -> dict:
    """Extract text (and optionally images) from PDF bytes."""
    from pypdf import PdfReader

    result = {"text": "", "images": [], "error": None}
    reader = PdfReader(io.BytesIO(content))
    pages_text = []

    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)

        # Attempt image extraction
        try:
            for img_obj in page.images:
                img_path = _save_image(img_obj.data, img_obj.name or "img")
                if img_path:
                    result["images"].append(img_path)
        except Exception:
            pass

    result["text"] = "\n".join(pages_text)
    return result


def _extract_docx(content: bytes) -> dict:
    """Extract text and images from DOCX bytes."""
    from docx import Document
    from docx.oxml.ns import qn

    result = {"text": "", "images": [], "error": None}
    doc = Document(io.BytesIO(content))
    paragraphs = []

    for para in doc.paragraphs:
        paragraphs.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(row_data))

    # Images
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                img_name = rel.target_ref.split("/")[-1]
                img_path = _save_image(img_data, img_name)
                if img_path:
                    result["images"].append(img_path)
            except Exception:
                pass

    result["text"] = "\n".join(paragraphs)
    return result


def _save_image(data: bytes, name: str) -> str | None:
    """Save image bytes to disk and return the path."""
    try:
        from PIL import Image

        ext = os.path.splitext(name)[-1].lower() or ".png"
        if ext not in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"):
            ext = ".png"
        uid = uuid.uuid4().hex[:8]
        filename = f"{uid}{ext}"
        path = os.path.join(IMAGES_DIR, filename)
        img = Image.open(io.BytesIO(data))
        img.save(path)
        return path
    except Exception as e:
        logger.warning(f"Image save failed: {e}")
        return None
